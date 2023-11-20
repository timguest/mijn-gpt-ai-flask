import requests
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup
from typing import Set, Optional, List, Any
from os import mkdir
import json
from pathlib import Path
from docx import Document
import time

SCRAPED_DATA = Path(__file__).parent / "scraped_data"


class Scraper:
    def __init__(self, url: str = ""):
        self.url = url
        self.visited = set()

    @property
    def url(self) -> str:
        return self._url

    @url.setter
    def url(self, url: url) -> None:
        self._url = url
        self.visited = set()  # Reset visited when URL is set

    def is_valid_url(self, url: str) -> bool:
        """Check if the URL is valid and within the same domain as the base URL."""
        parsed = urlparse(url)
        return bool(parsed.netloc) and parsed.netloc == urlparse(self.url).netloc

    def clean_url(self, href: str) -> str:
        """Strips irrelevant data from the url"""
        href = urljoin(self.url, href)
        parsed_href = urlparse(href)
        return parsed_href.scheme + "://" + parsed_href.netloc + parsed_href.path

    def get_domain(self, url: str) -> str:
        """Extracts the second-level domain from the URL."""
        return urlparse(url).netloc.split('.')[0]

    def get_all_page_links(self, url) -> list[Any]:
        """Returns all unique internal links found on a single page `url`."""
        try:
            response = requests.get(url, timeout=5)
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                links = [a.get('href') for a in soup.find_all('a', href=True)]
                filtered_links = [link for link in links if url in link]
                unique_filtered_links = list(set(filtered_links))
                return unique_filtered_links
        except requests.RequestException as e:
            print(f'Error accessing {url}: {e}')


    def create_storage_dir(self) -> Path:
        """Creates a directory within scraped_data/ named after the website's second-level domain."""
        domain = self.get_domain(self.url)
        dir_name = SCRAPED_DATA / domain

        # Create the directory and any missing parents
        dir_name.mkdir(parents=True, exist_ok=True)

        return dir_name

    def create_file_name(self, extension: str) -> Path:
        """Generates a filename called 'data' with the appropriate extension within the domain directory."""
        storage_dir = self.create_storage_dir()
        file_with_extension = f"data.{extension}"
        return storage_dir / file_with_extension

    def parse_response(self, response: requests.Response) -> None:
        """This method takes the response, uses soup to parse it and then writes everything to a file."""
        soup = BeautifulSoup(response.text, "html.parser")
        paragraphs = soup.find_all("p")

        if not paragraphs:
            print(f"no data for {response.url}")
            return

        for paragraph in paragraphs:
            content = paragraph.get_text()
            file_name = self.create_file_name(response)
            self.write_to_file(file_name, content)

    def scrape(self, url: Optional[str] = None) -> None:
        """
        Starts the scraping.

        A URL is required for this to work.

        The URL can be entered when instantiating the scraper:
        scraper = Scraper(url)

        or manually:
        scraper = Scraper()
        scraper.url = url

        or when calling this method:
        scraper = Scraper()
        scraper.scrape(url)

        Results are APPENDED to files in the scraped_data directory.


        """
        if url:
            self.url = url
        print(self.url)
        self.storage_dir = self.create_storage_dir()
        routes = self.get_all_page_links(url)
        print(routes)

        self.create_word_document_and_text_file(routes)

    def write_to_file(self, file_name: str, content: str):
        file_name = Path(file_name)
        if file_name.suffix == ".json":
            json.dump(content, f)
            return

        with open(file_name, "a", encoding="utf-8") as f:
            f.write(str(content))

    def create_word_document_and_text_file(self, internal_links):
        print('Starting to create Word document and text file')

        # Create the file paths
        text_file_path = self.create_file_name('txt')
        word_document_path = self.create_file_name('docx')
        print(f"Text file will be saved as: {text_file_path}")
        print(f"Word document will be saved as: {word_document_path}")

        document = Document()  # Create a Word document instance
        all_content = []  # List to keep all contents for the text file

        # Scrape the main URL
        content = get_page_content(self.url)
        if content:
            content = sanitize_content(content)  # Sanitize content
            document.add_heading(self.url, level=1)
            document.add_paragraph(content)
            all_content.append(self.url + '\n' + content)  # Add main URL content to the list

        # Scrape the internal URLs
        for url in internal_links:
            content = get_page_content(url)
            if content:
                content = sanitize_content(content)  # Sanitize content

                # Word Document
                document.add_page_break()
                title = url.replace(self.url, '').strip('/')
                document.add_heading(title, level=1)

                cleaned_content = clean_text(content)
                try:
                    document.add_paragraph(cleaned_content)
                except:
                    print(title)
                    continue

                # Text File
                all_content.append(title + '\n' + content)  # Add internal URL content to the list
                time.sleep(1)  # Sleep to be gentle on the server

        # Save the Word document
        document.save(word_document_path)

        # Write all contents to the text file, separated by double newlines
        with open(text_file_path, 'w', encoding='utf-8') as text_file:
            text_file.write('\n\n'.join(all_content))

        print('done')


# Example
# scraper = Scraper()
# scraper.url = "https://www.scrapethissite.com/"
# scraper.scrape()
def clean_text(text):
    # Replace or remove non-XML compatible characters
    cleaned_text = text.replace('\x00', '').replace('\r', '').replace('\n', ' ')
    # Further cleaning as needed
    return cleaned_text

def get_internal_links(url, visited=None):
    if visited is None:
        visited = set()

    normalized_url = url.lower()
    if normalized_url in visited:
        return visited
    else:
        visited.add(normalized_url)

    # Check if 'blog' is in the path of the URL
    if 'blog' in urlparse(normalized_url).path:
        print('Skipping blog:', normalized_url)
        return visited

    try:
        response = requests.get(url, timeout=5)
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            links = [a.get('href') for a in soup.find_all('a', href=True)]
            for link in links:
                joined_url = urljoin(url, link)
                # Normalize the joined URL for consistency
                normalized_joined_url = joined_url.lower()
                if 'blog' not in urlparse(normalized_joined_url).path:
                    if urlparse(normalized_joined_url).netloc == urlparse(url).netloc:
                        # Recursively get internal links only if not already visited
                        if normalized_joined_url not in visited:
                            get_internal_links(normalized_joined_url, visited)
    except requests.RequestException as e:
        print(f'Error accessing {url}: {e}')
    finally:
        return visited



def get_page_content(url):
    try:
        response = requests.get(url)
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            texts = soup.stripped_strings
            content = ' '.join(texts)
            return content
        else:
            return None
    except requests.RequestException:
        return None


def sanitize_content(text):
    # Replace double newlines with a single newline
    return text.replace('\n\n', '\n')


# Example
# scraper = Scraper()
# scraper.url = "https://www.scrapethissite.com/"
# scraper.scrape()
